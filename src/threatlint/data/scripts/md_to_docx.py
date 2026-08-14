#!/usr/bin/env python3
"""Convert a markdown threat model report to a formatted Word document.

Usage: md_to_docx.py <input.md> <output.docx>
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Heading colour palette (dark teal for H1/H2, slate for H3/H4)
# ---------------------------------------------------------------------------
H1_COLOR = RGBColor(0x1F, 0x49, 0x7D)   # deep navy
H2_COLOR = RGBColor(0x2E, 0x74, 0xB5)   # mid blue
H3_COLOR = RGBColor(0x00, 0x70, 0x4C)   # dark green
H4_COLOR = RGBColor(0x37, 0x37, 0x37)   # near-black


def _set_heading_color(paragraph, color):
    for run in paragraph.runs:
        run.font.color.rgb = color


def _apply_inline(para, text):
    """Add runs to *para* with **bold** and `code` formatting applied."""
    parts = re.split(r'(\*\*[^*\n]+\*\*|`[^`\n]+`|\*[^*\n]+\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _add_table(doc, lines):
    """Parse a contiguous block of markdown table lines and add to doc."""
    rows = []
    for line in lines:
        # Skip the separator row (e.g., |---|---|)
        if re.match(r"^\s*\|[-:\s|]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.cell(ri, ci)
            # Clear default paragraph and rebuild with inline formatting
            cell.paragraphs[0].clear()
            _apply_inline(cell.paragraphs[0], cell_text)
            if ri == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()  # breathing room after table


def _add_code_block(doc, lines):
    if not lines:
        return
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def convert(md_text: str, output_path: str) -> None:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    lines = md_text.splitlines()
    i = 0
    table_buf: list[str] = []
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ---- code fence ----
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                _add_code_block(doc, code_buf)
                in_code = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ---- flush table buffer when non-table line is encountered ----
        if not line.strip().startswith("|") and table_buf:
            _add_table(doc, table_buf)
            table_buf = []

        # ---- table row ----
        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue

        # ---- horizontal rule ----
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            # Draw a thin bottom border on the paragraph
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ---- headings ----
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            p = doc.add_paragraph(style=style_map[level])
            _apply_inline(p, text)
            color_map = {1: H1_COLOR, 2: H2_COLOR, 3: H3_COLOR, 4: H4_COLOR}
            _set_heading_color(p, color_map[level])
            i += 1
            continue

        # ---- unordered list ----
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            depth = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, m.group(2))
            p.paragraph_format.left_indent = Inches(0.25 * (depth + 1))
            i += 1
            continue

        # ---- ordered list ----
        m = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _apply_inline(p, m.group(2))
            i += 1
            continue

        # ---- blank line ----
        if not line.strip():
            i += 1
            continue

        # ---- normal paragraph ----
        p = doc.add_paragraph()
        _apply_inline(p, line)
        i += 1

    # Flush any remaining blocks
    if table_buf:
        _add_table(doc, table_buf)
    if in_code and code_buf:
        _add_code_block(doc, code_buf)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)

    src = Path(sys.argv[1])
    if not src.exists():
        print(f"Input file not found: {src}", file=sys.stderr)
        sys.exit(1)

    convert(src.read_text(encoding="utf-8"), sys.argv[2])
